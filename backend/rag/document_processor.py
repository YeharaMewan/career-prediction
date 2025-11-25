"""
Document Processor

Handles multi-format document loading, text extraction, chunking, and metadata management.

Supported formats:
- PDF (.pdf)
- Microsoft Word (.docx)
- Text (.txt)
- Markdown (.md)
- JSON (.json)
- YAML (.yaml, .yml)
"""

import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import hashlib
import json
import yaml

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_core.documents import Document

# Optional: python-docx for Word document support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None
    logger = logging.getLogger(__name__)
    logger.warning("python-docx not installed. DOCX file support disabled.")

from .config import (
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    SEPARATORS,
    SUPPORTED_PDF_EXTENSIONS,
    MAX_PDF_SIZE_MB,
)

# Import tracing utilities
try:
    from ..utils.tracing import (
        trace_sync,
        trace_async,
        add_span_attribute,
        add_span_attributes,
        add_span_event
    )
except ImportError:
    # Graceful fallback if tracing not available
    def trace_sync(*args, **kwargs):
        def decorator(func):
            return func
        return decorator

    def trace_async(*args, **kwargs):
        def decorator(func):
            return func
        return decorator

    def add_span_attribute(*args, **kwargs):
        pass

    def add_span_attributes(*args, **kwargs):
        pass

    def add_span_event(*args, **kwargs):
        pass

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# Supported file extensions by type
SUPPORTED_EXTENSIONS = {
    'pdf': ['.pdf'],
    'docx': ['.docx', '.doc'],
    'text': ['.txt'],
    'markdown': ['.md', '.markdown'],
    'json': ['.json'],
    'yaml': ['.yaml', '.yml']
}

# All supported extensions flattened
ALL_SUPPORTED_EXTENSIONS = [ext for exts in SUPPORTED_EXTENSIONS.values() for ext in exts]

# Maximum file size per format (in MB)
MAX_FILE_SIZES = {
    'pdf': MAX_PDF_SIZE_MB,
    'docx': 50,
    'text': 10,
    'markdown': 10,
    'json': 20,
    'yaml': 20
}


class DocumentProcessor:
    """
    Processes multi-format documents for RAG system:
    - Supports PDF, DOCX, TXT, MD, JSON, YAML
    - Extracts text with rich metadata
    - Chunks text for optimal embedding
    - Handles structured data (JSON/YAML) with flattening
    - Adds comprehensive metadata for tracking
    """

    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP,
        separators: List[str] = None,
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Overlap between chunks for context
            separators: Text splitting separators (hierarchical)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or SEPARATORS

        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=self.separators,
            length_function=len,
        )

        logger.info(
            f"DocumentProcessor initialized: chunk_size={chunk_size}, overlap={chunk_overlap}"
        )

    @trace_sync("document_processor.load_pdf")
    def load_pdf(self, pdf_path: Path, collection_type: str) -> List[Document]:
        """
        Load a single PDF file and extract text with metadata.

        Args:
            pdf_path: Path to PDF file
            collection_type: "academic" or "skill"

        Returns:
            List of LangChain Document objects with page-level metadata

        Raises:
            FileNotFoundError: If PDF doesn't exist
            ValueError: If file is too large or invalid format
        """
        # Validate file
        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF not found: {pdf_path}")

        if pdf_path.suffix.lower() not in SUPPORTED_PDF_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {pdf_path.suffix}")

        # Check file size
        file_size_mb = pdf_path.stat().st_size / (1024 * 1024)
        if file_size_mb > MAX_PDF_SIZE_MB:
            raise ValueError(
                f"PDF too large: {file_size_mb:.2f}MB (max: {MAX_PDF_SIZE_MB}MB)"
            )

        # Add span attributes
        add_span_attributes({
            "file.path": str(pdf_path),
            "file.name": pdf_path.name,
            "file.type": "pdf",
            "file.size_mb": round(file_size_mb, 2),
            "collection.name": collection_type
        })

        logger.info(f"Loading PDF: {pdf_path.name} ({file_size_mb:.2f}MB)")

        add_span_event("loading_pdf_with_pymupdf")

        # Load PDF with PyMuPDF
        loader = PyMuPDFLoader(str(pdf_path))
        documents = loader.load()

        # Add collection type to metadata
        for doc in documents:
            doc.metadata["collection_type"] = collection_type
            doc.metadata["source_file"] = pdf_path.name
            doc.metadata["file_path"] = str(pdf_path)

        add_span_attributes({
            "pages_loaded": len(documents),
            "documents_count": len(documents)
        })

        add_span_event("pdf_loaded_successfully", {
            "pages": len(documents)
        })

        logger.info(f"Loaded {len(documents)} pages from {pdf_path.name}")
        return documents

    def load_directory(
        self, directory: Path, collection_type: str
    ) -> List[Document]:
        """
        Load all PDFs from a directory.

        Args:
            directory: Directory containing PDFs
            collection_type: "academic" or "skill"

        Returns:
            List of all documents from all PDFs
        """
        if not directory.exists():
            logger.warning(f"Directory not found: {directory}")
            return []

        pdf_files = list(directory.glob("*.pdf"))
        if not pdf_files:
            logger.warning(f"No PDF files found in {directory}")
            return []

        logger.info(f"Found {len(pdf_files)} PDF files in {directory}")

        all_documents = []
        for pdf_file in pdf_files:
            try:
                docs = self.load_pdf(pdf_file, collection_type)
                all_documents.extend(docs)
            except Exception as e:
                logger.error(f"Error loading {pdf_file.name}: {e}")
                continue

        logger.info(f"Total documents loaded: {len(all_documents)}")
        return all_documents

    @staticmethod
    def detect_file_type(file_path: Path) -> Optional[str]:
        """
        Detect file type based on extension.

        Args:
            file_path: Path to file

        Returns:
            File type string or None if unsupported
        """
        suffix = file_path.suffix.lower()
        for file_type, extensions in SUPPORTED_EXTENSIONS.items():
            if suffix in extensions:
                return file_type
        return None

    @staticmethod
    def calculate_file_hash(file_path: Path) -> str:
        """
        Calculate SHA-256 hash of file content for change detection.

        Args:
            file_path: Path to file

        Returns:
            SHA-256 hash as hex string
        """
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            # Read in chunks to handle large files
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    def _validate_file(self, file_path: Path, file_type: str):
        """
        Validate file existence and size.

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is too large or invalid
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        max_size = MAX_FILE_SIZES.get(file_type, 50)

        if file_size_mb > max_size:
            raise ValueError(
                f"File too large: {file_size_mb:.2f}MB (max: {max_size}MB for {file_type})"
            )

    @trace_sync("document_processor.load_docx")
    def load_docx(self, docx_path: Path, collection_type: str) -> List[Document]:
        """
        Load a DOCX file and extract text.

        Args:
            docx_path: Path to DOCX file
            collection_type: Collection type identifier

        Returns:
            List of Document objects (one per paragraph or logical section)
        """
        # Check if python-docx is available
        if not DOCX_AVAILABLE:
            logger.warning(f"Skipping {docx_path.name}: python-docx not installed. Install with: pip install python-docx")
            return []

        self._validate_file(docx_path, 'docx')

        add_span_attributes({
            "file.path": str(docx_path),
            "file.name": docx_path.name,
            "file.type": "docx",
            "file.size_mb": round(docx_path.stat().st_size / (1024 * 1024), 2),
            "collection.name": collection_type
        })

        logger.info(f"Loading DOCX: {docx_path.name}")

        doc = DocxDocument(str(docx_path))
        documents = []

        # Extract text from paragraphs
        full_text = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:  # Skip empty paragraphs
                full_text.append(text)

        # Create a single document with all content
        if full_text:
            content = "\n\n".join(full_text)
            file_hash = self.calculate_file_hash(docx_path)

            doc_obj = Document(
                page_content=content,
                metadata={
                    "source": str(docx_path),
                    "source_file": docx_path.name,
                    "file_path": str(docx_path),
                    "file_type": "docx",
                    "collection_type": collection_type,
                    "file_hash": file_hash,
                    "paragraph_count": len(full_text),
                    "loaded_at": datetime.now().isoformat()
                }
            )
            documents.append(doc_obj)

        logger.info(f"Loaded DOCX with {len(full_text)} paragraphs")
        return documents

    def load_text(self, text_path: Path, collection_type: str) -> List[Document]:
        """
        Load a plain text file.

        Args:
            text_path: Path to text file
            collection_type: Collection type identifier

        Returns:
            List with single Document object
        """
        self._validate_file(text_path, 'text')

        logger.info(f"Loading text file: {text_path.name}")

        with open(text_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        if not content.strip():
            logger.warning(f"Empty text file: {text_path}")
            return []

        file_hash = self.calculate_file_hash(text_path)

        doc = Document(
            page_content=content,
            metadata={
                "source": str(text_path),
                "source_file": text_path.name,
                "file_path": str(text_path),
                "file_type": "text",
                "collection_type": collection_type,
                "file_hash": file_hash,
                "char_count": len(content),
                "loaded_at": datetime.now().isoformat()
            }
        )

        logger.info(f"Loaded text file ({len(content)} characters)")
        return [doc]

    def load_markdown(self, md_path: Path, collection_type: str) -> List[Document]:
        """
        Load a Markdown file.

        Args:
            md_path: Path to Markdown file
            collection_type: Collection type identifier

        Returns:
            List with single Document object
        """
        self._validate_file(md_path, 'markdown')

        logger.info(f"Loading markdown file: {md_path.name}")

        with open(md_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        if not content.strip():
            logger.warning(f"Empty markdown file: {md_path}")
            return []

        file_hash = self.calculate_file_hash(md_path)

        doc = Document(
            page_content=content,
            metadata={
                "source": str(md_path),
                "source_file": md_path.name,
                "file_path": str(md_path),
                "file_type": "markdown",
                "collection_type": collection_type,
                "file_hash": file_hash,
                "char_count": len(content),
                "loaded_at": datetime.now().isoformat()
            }
        )

        logger.info(f"Loaded markdown file ({len(content)} characters)")
        return [doc]

    def _flatten_json(
        self,
        data: Any,
        parent_key: str = '',
        sep: str = '.'
    ) -> Dict[str, Any]:
        """
        Flatten nested JSON/dict structure.

        Args:
            data: Data to flatten (dict, list, or primitive)
            parent_key: Parent key for nested items
            sep: Separator for nested keys

        Returns:
            Flattened dictionary
        """
        items = []

        if isinstance(data, dict):
            for k, v in data.items():
                new_key = f"{parent_key}{sep}{k}" if parent_key else k
                if isinstance(v, (dict, list)):
                    items.extend(self._flatten_json(v, new_key, sep=sep).items())
                else:
                    items.append((new_key, v))
        elif isinstance(data, list):
            for i, v in enumerate(data):
                new_key = f"{parent_key}[{i}]"
                if isinstance(v, (dict, list)):
                    items.extend(self._flatten_json(v, new_key, sep=sep).items())
                else:
                    items.append((new_key, v))
        else:
            items.append((parent_key, data))

        return dict(items)

    def _structured_to_text(self, data: Dict[str, Any], title: str = "") -> str:
        """
        Convert structured data to human-readable text for embedding.

        Args:
            data: Flattened data dictionary
            title: Optional title for the document

        Returns:
            Formatted text representation
        """
        lines = []

        if title:
            lines.append(f"# {title}")
            lines.append("")

        # Group by top-level keys
        grouped = {}
        for key, value in data.items():
            top_key = key.split('.')[0].split('[')[0]
            if top_key not in grouped:
                grouped[top_key] = []
            grouped[top_key].append((key, value))

        # Format each group
        for top_key, items in grouped.items():
            lines.append(f"## {top_key.replace('_', ' ').title()}")
            for key, value in items:
                # Clean up key for display
                display_key = key.replace('_', ' ').replace('.', ' > ')
                lines.append(f"- {display_key}: {value}")
            lines.append("")

        return "\n".join(lines)

    def load_json(self, json_path: Path, collection_type: str) -> List[Document]:
        """
        Load a JSON file and convert to searchable text.

        Args:
            json_path: Path to JSON file
            collection_type: Collection type identifier

        Returns:
            List of Document objects
        """
        self._validate_file(json_path, 'json')

        logger.info(f"Loading JSON file: {json_path.name}")

        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Flatten nested structure
        flattened = self._flatten_json(data)

        # Convert to readable text
        text_content = self._structured_to_text(
            flattened,
            title=json_path.stem.replace('_', ' ').title()
        )

        file_hash = self.calculate_file_hash(json_path)

        # Also keep original JSON as metadata
        doc = Document(
            page_content=text_content,
            metadata={
                "source": str(json_path),
                "source_file": json_path.name,
                "file_path": str(json_path),
                "file_type": "json",
                "collection_type": collection_type,
                "file_hash": file_hash,
                "original_structure": json.dumps(data)[:1000],  # First 1000 chars
                "key_count": len(flattened),
                "loaded_at": datetime.now().isoformat()
            }
        )

        logger.info(f"Loaded JSON file ({len(flattened)} keys)")
        return [doc]

    def load_yaml(self, yaml_path: Path, collection_type: str) -> List[Document]:
        """
        Load a YAML file and convert to searchable text.

        Args:
            yaml_path: Path to YAML file
            collection_type: Collection type identifier

        Returns:
            List of Document objects
        """
        self._validate_file(yaml_path, 'yaml')

        logger.info(f"Loading YAML file: {yaml_path.name}")

        with open(yaml_path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)

        # Flatten nested structure
        flattened = self._flatten_json(data)

        # Convert to readable text
        text_content = self._structured_to_text(
            flattened,
            title=yaml_path.stem.replace('_', ' ').title()
        )

        file_hash = self.calculate_file_hash(yaml_path)

        doc = Document(
            page_content=text_content,
            metadata={
                "source": str(yaml_path),
                "source_file": yaml_path.name,
                "file_path": str(yaml_path),
                "file_type": "yaml",
                "collection_type": collection_type,
                "file_hash": file_hash,
                "original_structure": json.dumps(data)[:1000],  # First 1000 chars
                "key_count": len(flattened),
                "loaded_at": datetime.now().isoformat()
            }
        )

        logger.info(f"Loaded YAML file ({len(flattened)} keys)")
        return [doc]

    def load_document(
        self,
        file_path: Path,
        collection_type: str
    ) -> List[Document]:
        """
        Universal document loader - detects type and uses appropriate loader.

        Args:
            file_path: Path to document file
            collection_type: Collection type identifier

        Returns:
            List of Document objects

        Raises:
            ValueError: If file type is unsupported
        """
        file_type = self.detect_file_type(file_path)

        if file_type is None:
            raise ValueError(
                f"Unsupported file type: {file_path.suffix}. "
                f"Supported: {ALL_SUPPORTED_EXTENSIONS}"
            )

        # Route to appropriate loader
        loaders = {
            'pdf': self.load_pdf,
            'docx': self.load_docx,
            'text': self.load_text,
            'markdown': self.load_markdown,
            'json': self.load_json,
            'yaml': self.load_yaml
        }

        loader = loaders[file_type]
        return loader(file_path, collection_type)

    def chunk_documents(
        self, documents: List[Document], collection_type: str
    ) -> List[Document]:
        """
        Split documents into chunks with enhanced metadata.

        Args:
            documents: List of documents to chunk
            collection_type: "academic" or "skill"

        Returns:
            List of chunked documents with metadata
        """
        logger.info(f"Chunking {len(documents)} documents...")

        # Split documents
        chunked_docs = self.text_splitter.split_documents(documents)

        # Add enhanced metadata to chunks
        timestamp = datetime.now().isoformat()
        for idx, chunk in enumerate(chunked_docs):
            # Generate unique chunk ID
            chunk_id = self._generate_chunk_id(chunk.page_content, idx)

            # Add metadata
            chunk.metadata.update(
                {
                    "chunk_id": chunk_id,
                    "chunk_index": idx,
                    "total_chunks": len(chunked_docs),
                    "timestamp": timestamp,
                    "collection_type": collection_type,
                    "chunk_length": len(chunk.page_content),
                }
            )

        logger.info(f"Created {len(chunked_docs)} chunks")
        return chunked_docs

    def process_file(
        self, file_path: Path, collection_type: str
    ) -> List[Document]:
        """
        Complete processing pipeline for a single file (any supported format).

        Args:
            file_path: Path to document file
            collection_type: Collection type identifier

        Returns:
            List of chunked documents ready for embedding
        """
        # Load document (auto-detects format)
        documents = self.load_document(file_path, collection_type)

        if not documents:
            return []

        # Chunk documents
        chunked_docs = self.chunk_documents(documents, collection_type)

        return chunked_docs

    def process_pdf(
        self, pdf_path: Path, collection_type: str
    ) -> List[Document]:
        """
        Complete processing pipeline for a single PDF.
        (Kept for backwards compatibility, use process_file instead)

        Args:
            pdf_path: Path to PDF file
            collection_type: "academic" or "skill"

        Returns:
            List of chunked documents ready for embedding
        """
        return self.process_file(pdf_path, collection_type)

    def process_directory(
        self, directory: Path, collection_type: str, file_patterns: List[str] = None
    ) -> List[Document]:
        """
        Complete processing pipeline for all supported files in a directory.

        Args:
            directory: Directory containing documents
            collection_type: Collection type identifier
            file_patterns: Optional list of glob patterns (e.g., ["*.pdf", "*.json"])
                          If None, processes all supported file types

        Returns:
            List of all chunked documents ready for embedding
        """
        if not directory.exists():
            logger.warning(f"Directory not found: {directory}")
            return []

        # Determine which files to process
        if file_patterns is None:
            file_patterns = [f"*{ext}" for ext in ALL_SUPPORTED_EXTENSIONS]

        # Find all matching files
        all_files = []
        for pattern in file_patterns:
            all_files.extend(directory.glob(pattern))

        # Remove duplicates
        all_files = list(set(all_files))

        if not all_files:
            logger.warning(f"No supported files found in {directory}")
            return []

        logger.info(f"Found {len(all_files)} files in {directory}")

        # Load all documents
        all_documents = []
        for file_path in all_files:
            try:
                docs = self.load_document(file_path, collection_type)
                all_documents.extend(docs)
            except Exception as e:
                logger.error(f"Error loading {file_path.name}: {e}")
                continue

        if not all_documents:
            return []

        logger.info(f"Total documents loaded: {len(all_documents)}")

        # Chunk documents
        chunked_docs = self.chunk_documents(all_documents, collection_type)

        return chunked_docs

    @staticmethod
    def _generate_chunk_id(content: str, index: int) -> str:
        """
        Generate a unique ID for a chunk based on content hash.

        Args:
            content: Chunk content
            index: Chunk index

        Returns:
            Unique chunk ID
        """
        content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
        return f"chunk_{index}_{content_hash}"

    def get_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """
        Get statistics about processed documents.

        Args:
            documents: List of documents

        Returns:
            Dictionary with statistics
        """
        if not documents:
            return {"total_chunks": 0, "total_chars": 0, "avg_chunk_size": 0}

        total_chars = sum(len(doc.page_content) for doc in documents)
        avg_chunk_size = total_chars / len(documents)

        # Get unique sources
        sources = set(doc.metadata.get("source_file", "unknown") for doc in documents)

        stats = {
            "total_chunks": len(documents),
            "total_chars": total_chars,
            "avg_chunk_size": int(avg_chunk_size),
            "unique_sources": len(sources),
            "sources": list(sources),
        }

        return stats


# Example usage
if __name__ == "__main__":
    from .config import ACADEMIC_DATA_DIR, SKILL_DATA_DIR

    processor = DocumentProcessor()

    # Process academic documents
    print("\n=== Processing Academic Documents ===")
    academic_docs = processor.process_directory(ACADEMIC_DATA_DIR, "academic")
    if academic_docs:
        stats = processor.get_stats(academic_docs)
        print(f"Academic documents: {stats}")

    # Process skill documents
    print("\n=== Processing Skill Documents ===")
    skill_docs = processor.process_directory(SKILL_DATA_DIR, "skill")
    if skill_docs:
        stats = processor.get_stats(skill_docs)
        print(f"Skill documents: {stats}")
